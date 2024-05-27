import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

# JSON data
data =  {
        "name": "Vishal Pachpande",
        "email": "vishalpachpande831@gmail.com",
        "phone": "+91 8668734615",
        "education": [
            {
                "institution": "Shri Sant Gadge Baba College of Engineering and Technology, Bhusawal",
                "degree": "B.Tech in Artificial Intelligence and Data Science",
                "years": "2021-2024"
            },
            {
                "institution": "Veermata Jijabai Technological Institute, Mumbai",
                "degree": "Diploma in Mechanical Engineering",
                "years": "2017-2020"
            }
        ],
        "experience": [
            {
                "company": "Hidevs Community",
                "role": "Gen Ai Intern",
                "duration": "May 2024 - Present",
                "responsibilities": [
                    "Developing applications powered by Large Language Models (LLM)",
                    "Utilizing Prompt Engineering for the development of applications"
                ]
            },
            {
                "company": "Ordnance Factory Bhusawal",
                "role": "Data Science Intern",
                "duration": "Sep 2023 - Nov 2023",
                "responsibilities": [
                    "Developed a computer vision application to detect text from rough and metal surfaces",
                    "Designed automation software to automate data entry into Tally Prime"
                ]
            },
            {
                "company": "Happymonk.ai, Bengaluru",
                "role": "Data Science Intern",
                "duration": "Sep 2022 - Dec 2022",
                "responsibilities": [
                    "Led a computer vision project to develop a model for detecting road cracks and potholes",
                    "Collected data for a computer vision-based security application to detect objects from security cameras"
                ]
            },
            {
                "company": "JSW Steel, Dolvi, Maharashtra",
                "role": "Junior Executive",
                "duration": "Nov 2020 – Jan 2022",
                "responsibilities": [
                    "Led a mechanical maintenance team of technicians to maintain Billet Caster Mechanical maintenance",
                    "Devised cleaning of flashback arrestor using ultrasonic cleaning method, resulting in 50% cost saving on flashback arrestors"
                ]
            }
        ],
        "projects": [
            {
                "title": "Text Extraction from Metal Parts and Rough Surfaces",
                "description": "Computer vision project to extract text from metal parts and rough surfaces"
            },
            {
                "title": "Fire Detection using Computer Vision",
                "description": "Computer vision application to detect industrial fires using cameras"
            }
        ],
        "skills": {
            "Computer Vision": "Completed projects",
            "Python": "Completed projects and proficient",
            "Machine Learning": "Completed projects",
            "Data Analysis": "Proficient",
            "Data Structures and Algorithms": "Intermediate"
        },
        "certifications": [
            "Huawei Certified ICT Associate in Cloud Computing",
            "Huawei Certified ICT Associate (HCIA) in AI Technology",
            "Joy of Computing using Python - IIT Madras",
            "Pregrad - Data Science Training"
        ]
    }



# Create a new Document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Add Name
name = doc.add_heading(data['name'], level=1)
name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add Contact Information
contact_info = doc.add_paragraph()
contact_info.add_run(f"Email: {data['email']} | Phone: {data['phone']}").bold = True
contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add Education Section
doc.add_heading("EDUCATION", level=2).bold = True
for edu in data['education']:
    edu_paragraph = doc.add_paragraph()
    edu_paragraph.add_run(f"{edu['institution']}\n").bold = True
    edu_paragraph.add_run(f"{edu['degree']} ({edu['years']})")
    edu_paragraph.space_after = Pt(6)

# Add Experience Section
doc.add_heading("EXPERIENCE", level=2).bold = True
for exp in data['experience']:
    exp_paragraph = doc.add_paragraph()
    exp_paragraph.add_run(f"{exp['company']} — {exp['role']}\n").bold = True
    exp_paragraph.add_run(f"{exp['duration']}")
    for responsibility in exp['responsibilities']:
        exp_paragraph.add_run(f"\n● {responsibility}")

# Add Projects Section
doc.add_heading("PROJECTS", level=2).bold = True
for project in data['projects']:
    project_paragraph = doc.add_paragraph()
    project_paragraph.add_run(f"{project['title']}\n").bold = True
    project_paragraph.add_run(project['description'])

# Add Skills Section
doc.add_heading("SKILLS", level=2).bold = True
skills_paragraph = doc.add_paragraph()
for skill, proficiency in data['skills'].items():
    skills_paragraph.add_run(f"{skill}: {proficiency}\n")

# Add Certifications Section
doc.add_heading("CERTIFICATIONS", level=2).bold = True
certs_paragraph = doc.add_paragraph()
for cert in data['certifications']:
    certs_paragraph.add_run(f"{cert}\n")

# Save the Document
doc.save("Vishal_Pachpande_Resume.docx")
